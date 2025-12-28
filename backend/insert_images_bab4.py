#!/usr/bin/env python3
"""
Script untuk menambahkan gambar ke dokumen BAB IV
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def insert_images_to_bab4():
    """Menambahkan gambar hasil test ke dokumen BAB IV"""
    
    # Load dokumen
    doc_path = 'backend/BAB_IV_HASIL_DAN_PEMBAHASAN.docx'
    doc = Document(doc_path)
    
    # Mapping gambar
    image_mappings = {
        '[Gambar 4.1 akan dimasukkan di sini: normal_driving_20251228_103910.jpg]': 
            'backend/test_results/normal_driving_20251228_103910.jpg',
        '[Gambar 4.2 akan dimasukkan di sini: simulated_drowsiness_20251228_104027.jpg]': 
            'backend/test_results/simulated_drowsiness_20251228_104027.jpg',
        '[Gambar 4.3 akan dimasukkan di sini: blinking_20251228_104133.jpg]': 
            'backend/test_results/blinking_20251228_104133.jpg',
        '[Gambar 4.4 akan dimasukkan di sini: low_light_20251228_104355.jpg]': 
            'backend/test_results/low_light_20251228_104355.jpg',
        '[Gambar 4.5 akan dimasukkan di sini: bright_light_20251228_104500.jpg]': 
            'backend/test_results/bright_light_20251228_104500.jpg',
    }
    
    # Cari dan replace placeholder dengan gambar
    for paragraph in doc.paragraphs:
        for placeholder, image_path in image_mappings.items():
            if placeholder in paragraph.text:
                # Hapus teks placeholder
                paragraph.text = ''
                
                # Tambahkan gambar
                if os.path.exists(image_path):
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(5.0))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"✅ Gambar ditambahkan: {image_path}")
                else:
                    paragraph.text = f"[ERROR: Gambar tidak ditemukan: {image_path}]"
                    print(f"❌ Gambar tidak ditemukan: {image_path}")
    
    # Save dokumen
    output_path = 'backend/BAB_IV_HASIL_DAN_PEMBAHASAN_FINAL.docx'
    doc.save(output_path)
    print(f"\n✅ Dokumen final berhasil dibuat: {output_path}")
    print(f"📄 Dokumen sudah lengkap dengan semua gambar!")

if __name__ == '__main__':
    insert_images_to_bab4()
